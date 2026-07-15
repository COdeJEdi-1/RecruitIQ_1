"""
Flask web app for the JD Generator — Arvind Limited.

Routes:
  GET  /                       → redirect to /generate if logged in, else /login
  GET  /login                  → login page (email/password + magic link)
  POST /api/auth/signin        → Supabase email+password sign-in
  POST /api/auth/magic-link    → Supabase magic link (passwordless)
  GET  /logout                 → clear session
  GET  /generate               → JD generator page (protected)
  GET  /history                → history page (protected) — sample JDs by department + user's approved JDs

  Legacy (Phase 1 — kept for backward compat):
  POST /api/generate           → single JD (seniority level based)
  POST /api/download           → download as .docx
  POST /api/approve            → save to KB (legacy)

  Arvind (Phase 2 — new endpoints):
  POST /api/jd/generate        → single JD at temperature 0.9 (no variants)
  POST /api/jd/download-html   → Arvind-branded HTML (open in browser, print→PDF)
  GET  /api/jd/preview         → live HTML preview in browser tab
  POST /api/jd/approve         → save to KB + log Signal 2 (edit diff)
  POST /api/jd/feedback        → log Signal 3 (explicit feedback)
  GET  /api/jd/download-pdf    → download as PDF (reportlab)
  GET  /api/jd/history         → user's approved JDs

  Sample JD Library (new — zero-LLM-cost browsing mode):
  GET  /library                → Sample JD Library page (protected)
  GET  /api/sample/filters     → division/department/family options for filter sidebar
  GET  /api/sample/list        → filtered list of sample JDs (dept/family/yoe_band)
  GET  /api/sample/get         → full JD content for preview modal (?ref=...)
  POST /api/sample/export-pdf  → export an "Edit a Copy" draft as PDF (never saved to KB)
  POST /api/sample/export-docx → export an "Edit a Copy" draft as DOCX (never saved to KB)
  POST /api/jd/share           → placeholder share-link endpoint (SSO wiring pending)
  GET  /shared/<token>         → placeholder shared-JD viewer (SSO wiring pending)
"""

import io
import json
import os
import re
import sys
import threading
import uuid
from datetime import datetime
from functools import wraps
from pathlib import Path

from dotenv import load_dotenv
from flask import (Flask, jsonify, make_response, redirect, render_template,
                   request, send_file, session, url_for)

load_dotenv()

sys.path.insert(0, os.path.dirname(__file__))
import scoring_service as scoring
from jd_parser import parse_jd_sections
from generate_jd import (
    DEPARTMENTS, DIVISIONS, SENIORITY_LEVELS, YOE_BANDS, YOE_LABELS,
    DEPARTMENT_CAPS_LABELS, DEPARTMENT_DISPLAY_NAMES, get_department_caps_groups,
    display_name, generate_jd, generate_jd_legacy, SENIORITY_TO_YOE,
)
from feedback_engine import (
    store_edit_diff,
    store_explicit_feedback,
    build_feedback_digest,
)
from community_skills import record_custom_skill, get_promoted_skills
from supabase_kb import upsert_sample_jd, SUPABASE_DB_ENABLED
import supabase_db as sdb
import darwinbox_service as darwin
from jd_constants import (JD_FOOTER_TEXT, JD_FOOTER_HEADING, JD_FOOTER_BRAND_LABEL,
                         JD_FOOTER_BRAND_NAME, append_jd_footer, strip_jd_footer)

# ── App setup ────────────────────────────────────────────────────────────────

app = Flask(__name__)
app.secret_key = os.getenv("FLASK_SECRET_KEY", "dev-secret-please-change-in-production")
app.jinja_env.filters["display_name"] = display_name


@app.context_processor
def inject_jd_constants():
    return {
        "jd_footer_text": JD_FOOTER_TEXT,
        "jd_footer_heading": JD_FOOTER_HEADING,
        "jd_footer_brand_label": JD_FOOTER_BRAND_LABEL,
        "jd_footer_brand_name": JD_FOOTER_BRAND_NAME,
    }

SUPABASE_URL = os.getenv("SUPABASE_URL")
SUPABASE_ANON_KEY = os.getenv("SUPABASE_ANON_KEY")
DEV_MODE = not SUPABASE_URL
ADMIN_EMAILS = {e.strip().lower() for e in os.getenv("ADMIN_EMAILS", "").split(",") if e.strip()}

if not DEV_MODE:
    from supabase import create_client
    supabase = create_client(SUPABASE_URL, SUPABASE_ANON_KEY)

# ── Microsoft SSO (Entra ID / Azure AD) ──────────────────────────────────────
MS_CLIENT_ID     = os.getenv("MICROSOFT_CLIENT_ID", "")
MS_CLIENT_SECRET = os.getenv("MICROSOFT_CLIENT_SECRET", "")
MS_TENANT_ID     = os.getenv("MICROSOFT_TENANT_ID", "common")
MS_REDIRECT_URI  = os.getenv("MICROSOFT_REDIRECT_URI", "http://localhost:5001/auth/microsoft/callback")
MS_AUTHORITY     = f"https://login.microsoftonline.com/{MS_TENANT_ID}"
MS_SCOPES        = ["User.Read"]
MS_ENABLED       = bool(MS_CLIENT_ID and MS_CLIENT_SECRET)

KB_ROOT = Path(__file__).parent / "kb"
STATIC_ROOT = Path(__file__).parent / "static"

# Pre-load Arvind logos as base64 data URIs (embedded in HTML/PDF, no external deps)
def _load_logo(name: str) -> str:
    p = STATIC_ROOT / name
    return p.read_text(encoding="utf-8").strip() if p.exists() else ""

LOGO_MAIN = _load_logo("arvind_logo_1_b64.txt")   # full Arvind Limited logo (header)
LOGO_GCC  = _load_logo("arvind_logo_2_b64.txt")   # compact GCC logo (footer)

# Pending generation results — persisted to disk so server restarts don't lose them
_PENDING_FILE = Path(__file__).parent / "darwin_data" / "pending_generations.json"


def _load_pending() -> dict:
    try:
        return json.loads(_PENDING_FILE.read_text(encoding="utf-8"))
    except Exception:
        return {}


def _save_pending(store: dict):
    try:
        _PENDING_FILE.parent.mkdir(exist_ok=True)
        _PENDING_FILE.write_text(json.dumps(store, default=str), encoding="utf-8")
    except Exception:
        pass


_pending_generations: dict = _load_pending()


# ── Auth helpers ──────────────────────────────────────────────────────────────

def _enrich_user(user: dict) -> dict:
    """Add computed fields (is_admin) to the session user dict."""
    user = dict(user)
    role = user.get("role", "user")
    user["is_admin"] = (role == "admin") or (not DEV_MODE and user.get("email", "").lower() in ADMIN_EMAILS)
    return user


def login_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if "user" not in session:
            # Return JSON 401 for API/XHR calls so the frontend can handle it
            if request.path.startswith("/api/") or request.is_json:
                return jsonify({"error": "Session expired. Please sign in again.", "redirect": url_for("login")}), 401
            return redirect(url_for("login"))
        return f(*args, **kwargs)
    return decorated


def admin_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if "user" not in session:
            return redirect(url_for("login"))
        user = _enrich_user(session["user"])
        if not user["is_admin"]:
            return redirect(url_for("library"))
        return f(*args, **kwargs)
    return decorated


def _make_initials(name: str) -> str:
    parts = name.split()
    return "".join(p[0].upper() for p in parts[:2]) if parts else "U"


# ── Page routes ───────────────────────────────────────────────────────────────

@app.route("/")
def index():
    return redirect(url_for("hub"))


@app.route("/hub")
def hub():
    """Platform landing page — no auth required. Entry point for both modules."""
    landing = Path(__file__).parent.parent / "index.html"
    if landing.exists():
        return landing.read_text(encoding="utf-8")
    return redirect(url_for("login"))


@app.route("/login")
def login():
    has_local_users = bool(_load_local_users())
    return render_template("login.html", dev_mode=DEV_MODE or has_local_users)


@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for("login"))


@app.route("/generate")
@login_required
def generator():
    family_labels = {fam: display_name(fam) for fams in DEPARTMENTS.values() for fam in fams}
    return render_template(
        "generator.html",
        user=_enrich_user(session["user"]),
        departments=DEPARTMENTS,
        family_labels=family_labels,
        divisions=DIVISIONS,
        department_caps_groups=get_department_caps_groups(),
        levels=SENIORITY_LEVELS,
        yoe_bands=YOE_BANDS,
        yoe_labels=YOE_LABELS,
        dev_mode=DEV_MODE,
    )


@app.route("/history")
@login_required
def history():
    return render_template(
        "history.html",
        user=_enrich_user(session["user"]),
        departments=DEPARTMENTS,
    )


@app.route("/admin")
@admin_required
def admin_panel():
    return render_template("admin.html", user=_enrich_user(session["user"]))


@app.route("/api/admin/stats")
@admin_required
def api_admin_stats():
    """Return aggregate metrics for the admin dashboard.
    Tries Supabase first (rich data); falls back to filesystem KB scan."""
    from datetime import datetime, timedelta, timezone

    sb_stats = sdb.get_admin_stats()
    if sb_stats is not None:
        # Merge filesystem dept display-name keys on top of raw keys
        sb_stats["dept_counts"] = {
            display_name(k): v for k, v in sb_stats["dept_counts"].items()
        }
        for item in sb_stats.get("recent_activity", []):
            item["dept"]   = display_name(item["dept"])
            item["family"] = display_name(item["family"])
        return jsonify(sb_stats)

    # ── Filesystem fallback ───────────────────────────────────────────────────
    now = datetime.now(timezone.utc)
    week_ago = (now - timedelta(days=7)).isoformat()
    two_weeks_ago = (now - timedelta(days=14)).isoformat()

    total = 0
    this_week = 0
    last_week = 0
    dept_counts: dict = {}
    recent_activity = []

    for dept_dir in sorted(KB_ROOT.iterdir()):
        if not dept_dir.is_dir():
            continue
        for family_dir in sorted(dept_dir.iterdir()):
            if not family_dir.is_dir():
                continue
            sample_dir = family_dir / "sample_jds"
            if not sample_dir.exists():
                continue
            for jd_file in sample_dir.iterdir():
                if not jd_file.name.endswith(".txt"):
                    continue
                raw = jd_file.read_text(encoding="utf-8").strip()
                try:
                    meta = json.loads(raw.split("\n")[0])
                except Exception:
                    continue
                ts = meta.get("approved_at") or meta.get("date_added", "")
                total += 1
                dept = dept_dir.name
                dept_counts[dept] = dept_counts.get(dept, 0) + 1
                if ts >= week_ago:
                    this_week += 1
                elif ts >= two_weeks_ago:
                    last_week += 1
                recent_activity.append({
                    "role": meta.get("role_title", "—"),
                    "dept": display_name(dept),
                    "family": display_name(family_dir.name),
                    "date": ts[:10] if ts else "—",
                    "by": meta.get("approved_by", "—"),
                })

    recent_activity.sort(key=lambda x: x["date"], reverse=True)

    # Mark the newest JD
    if recent_activity:
        recent_activity[0]["is_newest"] = True

    # Per-user breakdown: {email: {dept: count, total: count}}
    user_stats: dict = {}
    for item in recent_activity:
        by = item.get("by", "—")
        dept = item.get("dept", "—")
        if by not in user_stats:
            user_stats[by] = {"total": 0, "depts": {}}
        user_stats[by]["total"] += 1
        user_stats[by]["depts"][dept] = user_stats[by]["depts"].get(dept, 0) + 1

    # Daily counts — last 7 days
    from collections import defaultdict
    daily_counts: dict = defaultdict(int)
    for item in recent_activity:
        if item["date"] and item["date"] != "—":
            daily_counts[item["date"]] += 1
    # Sort and emit last 7 days
    all_days = sorted(daily_counts.keys())[-7:]
    daily_series = [{"date": d, "count": daily_counts[d]} for d in all_days]

    # AI alignment: avg overall_rating from feedback logs (0–5 → 0–100%)
    feedback_ratings = []
    for dept_dir in KB_ROOT.iterdir():
        if not dept_dir.is_dir(): continue
        for family_dir in dept_dir.iterdir():
            if not family_dir.is_dir(): continue
            fb_log = family_dir / "feedback" / "feedback_log.jsonl"
            if not fb_log.exists(): continue
            for line in fb_log.read_text(encoding="utf-8").splitlines():
                try:
                    fb = json.loads(line)
                    r = fb.get("overall_rating")
                    if isinstance(r, (int, float)) and 1 <= r <= 5:
                        feedback_ratings.append(r)
                except Exception:
                    pass
    ai_score = round(sum(feedback_ratings) / len(feedback_ratings) / 5 * 100, 1) if feedback_ratings else None

    dept_counts_display = {display_name(k): v for k, v in dept_counts.items()}
    return jsonify({
        "total_jds": total,
        "this_week": this_week,
        "last_week": last_week,
        "dept_counts": dept_counts_display,
        "recent_activity": recent_activity[:20],
        "user_stats": user_stats,
        "daily_series": daily_series,
        "ai_score": ai_score,
        "feedback_count": len(feedback_ratings),
    })


# ── Microsoft SSO routes ──────────────────────────────────────────────────────

@app.route("/auth/microsoft")
def auth_microsoft_start():
    """Redirect user to Microsoft login page."""
    if not MS_ENABLED:
        return redirect(url_for("login") + "?error=Microsoft+SSO+not+configured")
    import msal, secrets as _secrets
    state = _secrets.token_urlsafe(24)
    session["ms_oauth_state"] = state
    msal_app = msal.ConfidentialClientApplication(
        MS_CLIENT_ID,
        authority=MS_AUTHORITY,
        client_credential=MS_CLIENT_SECRET,
    )
    auth_url = msal_app.get_authorization_request_url(
        scopes=MS_SCOPES,
        state=state,
        redirect_uri=MS_REDIRECT_URI,
    )
    return redirect(auth_url)


@app.route("/auth/microsoft/callback")
def auth_microsoft_callback():
    """Handle Azure AD callback, exchange code for token, create session."""
    import msal, requests as _req
    error = request.args.get("error")
    if error:
        desc = request.args.get("error_description", error)
        return redirect(url_for("login") + f"?error={desc}")

    code  = request.args.get("code")
    state = request.args.get("state")
    if not code or state != session.pop("ms_oauth_state", None):
        return redirect(url_for("login") + "?error=Invalid+OAuth+state")

    msal_app = msal.ConfidentialClientApplication(
        MS_CLIENT_ID,
        authority=MS_AUTHORITY,
        client_credential=MS_CLIENT_SECRET,
    )
    token_result = msal_app.acquire_token_by_authorization_code(
        code=code,
        scopes=MS_SCOPES,
        redirect_uri=MS_REDIRECT_URI,
    )
    if "error" in token_result:
        desc = token_result.get("error_description", token_result["error"])
        return redirect(url_for("login") + f"?error={desc}")

    access_token = token_result["access_token"]
    headers = {"Authorization": f"Bearer {access_token}"}

    # Fetch user profile from Microsoft Graph
    profile = _req.get("https://graph.microsoft.com/v1.0/me", headers=headers).json()
    email = (profile.get("mail") or profile.get("userPrincipalName", "")).lower()
    name  = profile.get("displayName") or profile.get("givenName") or email.split("@")[0]

    # Try to fetch profile photo (small 48x48)
    photo_url = ""
    try:
        photo_resp = _req.get(
            "https://graph.microsoft.com/v1.0/me/photos/48x48/$value",
            headers=headers, timeout=4
        )
        if photo_resp.status_code == 200:
            import base64
            b64 = base64.b64encode(photo_resp.content).decode()
            ct  = photo_resp.headers.get("Content-Type", "image/jpeg")
            photo_url = f"data:{ct};base64,{b64}"
    except Exception:
        pass

    session["user"] = {
        "email":    email,
        "name":     name,
        "picture":  photo_url,
        "initials": _make_initials(name),
        "auth_provider": "microsoft",
    }
    return redirect(url_for("library"))


@app.route("/api/auth/microsoft/status")
def api_ms_status():
    """Return whether Microsoft SSO is configured (for the login page)."""
    return jsonify({"enabled": MS_ENABLED})


# ── Auth API endpoints ────────────────────────────────────────────────────────

def _load_local_users() -> list:
    """Load users from darwin_data/users.json for DEV_MODE auth."""
    uf = Path(__file__).parent / "darwin_data" / "users.json"
    try:
        return json.loads(uf.read_text(encoding="utf-8"))
    except Exception:
        return []


@app.route("/api/auth/signin", methods=["POST"])
def api_signin():
    from werkzeug.security import check_password_hash as _chk
    data = request.get_json(silent=True) or {}
    email = data.get("email", "").strip().lower()
    password = data.get("password", "")

    # Always try local users.json first (covers DEV_MODE and local recruiters
    # even when Supabase is configured).
    local_match = next((u for u in _load_local_users() if u.get("email", "").lower() == email), None)
    if local_match:
        if not _chk(local_match["password_hash"], password):
            return jsonify({"error": "Invalid email or password."}), 401
        name = local_match.get("full_name") or email.split("@")[0]
        session["user"] = {
            "user_id":  local_match["user_id"],
            "email":    local_match["email"],
            "name":     name,
            "picture":  "",
            "initials": _make_initials(name),
            "role":     local_match.get("role", "user"),
        }
        return jsonify({"redirect": url_for("library")})

    if DEV_MODE:
        return jsonify({"error": "Invalid email or password."}), 401

    data = request.get_json(silent=True) or {}
    email = data.get("email", "").strip()
    password = data.get("password", "").strip()

    if not email or not password:
        return jsonify({"error": "Email and password are required."}), 400

    mode = data.get("mode", "signin")
    try:
        if mode == "signup":
            resp = supabase.auth.sign_up({"email": email, "password": password})
            if resp.session is None:
                return jsonify({"message": "Account created! Check your email to confirm before signing in."})
            user = resp.user
        else:
            resp = supabase.auth.sign_in_with_password({"email": email, "password": password})
            user = resp.user

        name = (user.user_metadata or {}).get("full_name") or email.split("@")[0]
        session["user"] = {
            "email": user.email,
            "name": name,
            "picture": (user.user_metadata or {}).get("avatar_url", ""),
            "initials": _make_initials(name),
        }
        return jsonify({"redirect": url_for("library")})
    except Exception as exc:
        return jsonify({"error": str(exc)}), 401


@app.route("/api/auth/magic-link", methods=["POST"])
def api_magic_link():
    if DEV_MODE:
        session["user"] = {"email": "dev@arvind.in", "name": "Dev User",
                           "picture": "", "initials": "DU"}
        return jsonify({"message": "Dev mode — signed in directly."})

    data = request.get_json(silent=True) or {}
    email = data.get("email", "").strip()
    if not email:
        return jsonify({"error": "Email is required."}), 400

    try:
        callback_url = request.host_url.rstrip("/") + "/auth/callback"
        supabase.auth.sign_in_with_otp({"email": email, "options": {"email_redirect_to": callback_url}})
        return jsonify({"message": f"Magic link sent to {email}. Check your inbox."})
    except Exception as exc:
        return jsonify({"error": str(exc)}), 400


@app.route("/auth/callback")
def auth_callback():
    """Landing page for Supabase magic-link redirects. The token arrives in the
    URL fragment (#access_token=...) which is client-side only, so we serve a
    small HTML page that reads it and POSTs it to /api/auth/token-exchange."""
    return """<!DOCTYPE html>
<html><head><title>Signing you in…</title>
<style>body{font-family:sans-serif;display:flex;align-items:center;justify-content:center;
min-height:100vh;margin:0;background:#fff8f7;color:#251819;}
.msg{text-align:center;}.spinner{width:32px;height:32px;border:3px solid #fce2e2;
border-top-color:#810022;border-radius:50%;animation:spin .7s linear infinite;margin:0 auto 16px;}
@keyframes spin{to{transform:rotate(360deg)}}</style></head>
<body><div class="msg"><div class="spinner"></div><p>Completing sign-in…</p></div>
<script>
const hash = Object.fromEntries(new URLSearchParams(location.hash.slice(1)));
if (hash.access_token) {
  fetch('/api/auth/token-exchange', {
    method:'POST', headers:{'Content-Type':'application/json'},
    body: JSON.stringify({access_token: hash.access_token, refresh_token: hash.refresh_token})
  }).then(r => r.json()).then(d => {
    if (d.redirect) location.href = d.redirect;
    else document.querySelector('.msg').innerHTML = '<p style="color:#ba1a1a">Sign-in failed: ' + (d.error||'Unknown error') + '</p>';
  }).catch(() => { document.querySelector('.msg').innerHTML = '<p style="color:#ba1a1a">Sign-in failed. Please try again.</p>'; });
} else {
  document.querySelector('.msg').innerHTML = '<p style="color:#ba1a1a">Invalid or expired link. <a href="/login">Try again</a></p>';
}
</script></body></html>"""


@app.route("/api/auth/token-exchange", methods=["POST"])
def api_token_exchange():
    """Exchange a Supabase access token (from magic-link callback) for a Flask session."""
    data = request.get_json(silent=True) or {}
    access_token = data.get("access_token", "")
    refresh_token = data.get("refresh_token", "")
    if not access_token:
        return jsonify({"error": "No token provided."}), 400
    try:
        resp = supabase.auth.set_session(access_token, refresh_token)
        user = resp.user
        name = (user.user_metadata or {}).get("full_name") or user.email.split("@")[0]
        session["user"] = {
            "email": user.email,
            "name": name,
            "picture": (user.user_metadata or {}).get("avatar_url", ""),
            "initials": _make_initials(name),
        }
        return jsonify({"redirect": url_for("library")})
    except Exception as exc:
        return jsonify({"error": str(exc)}), 401


# ── Legacy API (Phase 1 — kept for backward compat) ──────────────────────────

@app.route("/api/generate", methods=["POST"])
@login_required
def api_generate():
    """Legacy single-JD endpoint. Maps seniority level → YoE band internally."""
    data = request.get_json(silent=True) or {}
    role_title = data.get("role_title", "").strip()
    department = data.get("department", "").strip()
    family = data.get("family", "").strip()
    level = data.get("level", "").strip()
    focus_areas = data.get("focus_areas", "").strip()

    if not all([role_title, department, family, level]):
        return jsonify({"error": "role_title, department, family, and level are all required."}), 400

    try:
        jd = generate_jd_legacy(
            role_title=role_title,
            department=department,
            family=family,
            seniority_level=level,
            focus_areas=focus_areas,
        )
        return jsonify({"jd": jd})
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/api/approve", methods=["POST"])
@login_required
def api_approve():
    """Legacy approve — saves to KB sample_jds without feedback tracking."""
    data = request.get_json(silent=True) or {}
    jd_text = data.get("jd", "").strip()
    role_title = data.get("role_title", "Unknown Role").strip()
    department = data.get("department", "").strip()
    family = data.get("family", "").strip()
    level = data.get("level", "").strip()

    if not jd_text:
        return jsonify({"error": "No JD content to save."}), 400

    yoe_band = SENIORITY_TO_YOE.get(level, "2-5")
    band_prefix = yoe_band.replace("+", "plus").replace("-", "_")

    sample_dir = KB_ROOT / department / family / "sample_jds"
    sample_dir.mkdir(parents=True, exist_ok=True)

    existing = sorted(sample_dir.glob(f"{band_prefix}_*.txt"))
    next_idx = len(existing) + 1
    filename = sample_dir / f"{band_prefix}_{next_idx:03d}.txt"

    metadata = {
        "role_family": family.replace("_", " ").title(),
        "yoe_band": yoe_band,
        "date_added": datetime.now().strftime("%Y-%m-%d"),
        "approved_by": session["user"]["email"],
    }
    filename.write_text(json.dumps(metadata) + "\n\n---\n\n" + append_jd_footer(jd_text), encoding="utf-8")

    return jsonify({
        "success": True,
        "message": f"JD saved to knowledge base as {filename.name}.",
        "file": str(filename.relative_to(Path(__file__).parent)),
    })


# ── Arvind Phase 2 API endpoints ──────────────────────────────────────────────

@app.route("/api/jd/generate", methods=["POST"])
@login_required
def api_jd_generate():
    """
    Generate exactly ONE JD at temperature 0.9 for the given role / YoE band.
    Injects feedback digest from previous approvals for this dept/family/band.
    """
    data = request.get_json(silent=True) or {}
    role_title = data.get("role_title", "").strip()
    department = data.get("department", "").strip()
    family = data.get("family", "").strip()
    yoe_band = data.get("yoe_band", "2-5").strip()
    focus_areas = data.get("focus_areas", "").strip()
    role_type        = data.get("role_type", "individual_contributor").strip()
    team_size        = data.get("team_size")  # int or None
    employment_type  = data.get("employment_type", "Full-Time").strip()
    work_mode        = data.get("work_mode", "Onsite").strip()
    reports_to       = data.get("reports_to", "").strip()
    must_have_skills = data.get("must_have_skills") or []  # list[str]
    custom_skills    = data.get("custom_skills") or []  # list[str] — user-typed, not auto-suggested

    if not all([role_title, department, family, yoe_band]):
        return jsonify({"error": "role_title, department, family, and yoe_band are required."}), 400

    if department not in DEPARTMENTS:
        return jsonify({"error": f"Unknown department '{department}'."}), 400

    if yoe_band not in YOE_BANDS:
        return jsonify({"error": f"Unknown yoe_band '{yoe_band}'. Choose from: {', '.join(YOE_BANDS)}"}), 400

    # Fetch feedback digest to inject into prompts
    feedback_digest = build_feedback_digest(department, family, yoe_band)
    if feedback_digest:
        print(f"\n[FEEDBACK DIGEST INJECTED] {department}/{family} · {yoe_band} yrs:\n{feedback_digest}\n", flush=True)
    else:
        print(f"\n[NO PRIOR FEEDBACK YET] {department}/{family} · {yoe_band} yrs — nothing to apply.\n", flush=True)

    try:
        jd_text = generate_jd(
            role_title=role_title,
            department=department,
            family=family,
            yoe_band=yoe_band,
            focus_areas=focus_areas,
            feedback_digest=feedback_digest,
            role_type=role_type,
            team_size=team_size,
            employment_type=employment_type,
            work_mode=work_mode,
            reports_to=reports_to,
            must_have_skills=must_have_skills if must_have_skills else None,
        )
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500

    generation_id = str(uuid.uuid4())

    # Cache the generated text so /api/jd/approve can compute an edit diff
    _pending_generations[generation_id] = {
        "role_title": role_title,
        "department": department,
        "family": family,
        "yoe_band": yoe_band,
        "focus_areas": focus_areas,
        "text": jd_text,
        "user_email": session["user"]["email"],
        "custom_skills": custom_skills,
    }
    _save_pending(_pending_generations)

    # Persist draft to Supabase (best-effort)
    sdb.save_draft(
        generation_id=generation_id,
        user_email=session["user"]["email"],
        role_title=role_title,
        department=department,
        division=_division_for_dept(department),
        family=family,
        yoe_band=yoe_band,
        focus_areas=focus_areas,
        original_text=jd_text,
    )

    return jsonify({"generation_id": generation_id, "jd": jd_text})


@app.route("/api/jd/approve", methods=["POST"])
@login_required
def api_jd_approve():
    """
    Save approved JD to KB sample_jds + log Signal 2 (edit diff).
    Returns edit_ratio so the UI can show it.
    """
    data = request.get_json(silent=True) or {}
    generation_id = data.get("generation_id", "").strip()
    final_jd = data.get("jd", "").strip()
    role_title = data.get("role_title", "").strip()
    department = data.get("department", "").strip()
    family = data.get("family", "").strip()
    yoe_band = data.get("yoe_band", "2-5").strip()

    if not final_jd:
        return jsonify({"error": "No JD content to save."}), 400

    # Original text from cache (to compute edit diff)
    gen = _pending_generations.get(generation_id, {})
    original_text = gen.get("text", "")

    # Save to KB
    band_prefix = yoe_band.replace("+", "plus").replace("-", "_")
    sample_dir = KB_ROOT / department / family / "sample_jds"
    sample_dir.mkdir(parents=True, exist_ok=True)

    jd_id = str(uuid.uuid4())[:8]
    existing = sorted(sample_dir.glob(f"{band_prefix}_*.txt"))
    next_idx = len(existing) + 1
    filename = sample_dir / f"{band_prefix}_{next_idx:03d}.txt"

    now = datetime.now()
    metadata = {
        "jd_id": jd_id,
        "role_family": family.replace("_", " ").title(),
        "department_caps": DEPARTMENT_CAPS_LABELS.get(department.lower(), department.upper()),
        "yoe_band": yoe_band,
        "generation_id": generation_id,
        "date_added": now.strftime("%Y-%m-%d"),
        "approved_at": now.isoformat(),
        "approved_by": session["user"]["email"],
    }
    final_for_diff = strip_jd_footer(final_jd)
    final_jd_to_save = append_jd_footer(final_jd)

    # Log community-added custom skills (best-effort — never blocks the save)
    try:
        for skill in gen.get("custom_skills", []):
            record_custom_skill(
                dept=department,
                family=family,
                skill_raw=skill,
                user_email=session["user"]["email"],
                jd_id=jd_id,
                role_title=role_title,
            )
    except Exception:
        pass

    # Log Signal 2 (body edits only — footer excluded from diff)
    edit_ratio = 0.0
    if original_text:
        edit_ratio = store_edit_diff(
            dept=department,
            family=family,
            yoe_band=yoe_band,
            jd_id=jd_id,
            generation_id=generation_id,
            original_text=original_text,
            final_text=final_for_diff,
            role_title=role_title,
            user_email=session["user"]["email"],
        )

    # Remove from pending cache
    _pending_generations.pop(generation_id, None)
    _save_pending(_pending_generations)

    # Best-effort Supabase persistence (never blocks filesystem-based flow)
    sdb.approve_draft(
        generation_id=generation_id,
        final_text=final_for_diff,
        jd_ref=jd_id,
        edit_ratio=edit_ratio,
    )
    sdb.save_edit_diff(
        jd_id=jd_id,
        generation_id=generation_id,
        user_email=session["user"]["email"],
        department=department,
        family=family,
        yoe_band=yoe_band,
        role_title=role_title,
        edit_ratio=edit_ratio,
        original_text=original_text,
        final_text=final_for_diff,
    )
    sdb.bump_jd_count(session["user"]["email"])
    sdb.bump_kb_version(department, family, yoe_band)

    if SUPABASE_DB_ENABLED:
        upsert_sample_jd(
            department=department,
            division=_division_for_dept(department),
            role_family=metadata["role_family"],
            yoe_band=yoe_band,
            seniority_label="",
            jd_text=final_jd_to_save,
            metadata=metadata,
            added_by=session["user"]["email"],
        )

    # Publish to Darwinbox — required for the JD to appear in DarwinBox portal
    darwin_error = None
    darwin_job = None
    try:
        darwin_job = darwin.publish_to_darwinbox(
            jd_id=jd_id,
            role_title=role_title or family.replace("_", " ").title(),
            jd_text=final_jd_to_save,
            created_by=session["user"].get("user_id"),
        )
    except Exception as exc:
        darwin_error = str(exc)

    return jsonify({
        "success": True,
        "jd_id": jd_id,
        "edit_ratio": edit_ratio,
        "message": "JD approved and saved to knowledge base.",
        "file": str(filename.relative_to(Path(__file__).parent)),
        "darwinbox_job_id": darwin_job["darwinbox_job_id"] if darwin_job else None,
        "darwin_error": darwin_error,
    })


@app.route("/api/jd/feedback", methods=["POST"])
@login_required
def api_jd_feedback():
    """Log Signal 3: explicit star rating + tags + free text."""
    data = request.get_json(silent=True) or {}

    required = ["jd_id", "department", "family", "yoe_band",
                "overall_rating", "role_title"]
    for field in required:
        if not data.get(field):
            return jsonify({"error": f"'{field}' is required."}), 400

    overall_rating = int(data["overall_rating"])
    if overall_rating not in range(1, 6):
        return jsonify({"error": "overall_rating must be 1–5."}), 400

    store_explicit_feedback(
        dept=data["department"],
        family=data["family"],
        yoe_band=data["yoe_band"],
        jd_id=data["jd_id"],
        overall_rating=overall_rating,
        section_ratings=data.get("section_ratings", {}),
        positive_tags=data.get("positive_tags", []),
        improvement_tags=data.get("improvement_tags", []),
        free_text=data.get("free_text", ""),
        better_than_manual=data.get("better_than_manual", "about_the_same"),
        role_title=data["role_title"],
        user_email=session["user"]["email"],
    )

    # Mirror feedback to Supabase (best-effort)
    sdb.save_feedback(
        jd_id=data["jd_id"],
        generation_id=data.get("generation_id", ""),
        user_email=session["user"]["email"],
        department=data["department"],
        family=data["family"],
        yoe_band=data["yoe_band"],
        role_title=data["role_title"],
        overall_rating=overall_rating,
        section_ratings=data.get("section_ratings", {}),
        positive_tags=data.get("positive_tags", []),
        improvement_tags=data.get("improvement_tags", []),
        free_text=data.get("free_text", ""),
        better_than_manual=data.get("better_than_manual", "about_the_same"),
    )

    return jsonify({"success": True, "message": "Feedback saved. Thank you!"})


def _build_export_context(data: dict) -> dict:
    """Parse JD text + merge metadata into a template context dict."""
    jd_text   = strip_jd_footer(data.get("jd", "").strip())
    sections  = parse_jd_sections(jd_text)
    role_title = data.get("role_title", "Job Description").strip()
    return {
        **sections,
        "role_title":       role_title,
        "department":       data.get("department", ""),
        "family":           data.get("family", ""),
        "division":         data.get("division", ""),
        "yoe_band":         data.get("yoe_band", ""),
        "location":         data.get("location", "Ahmedabad, Gujarat"),
        "employment_type":  data.get("employment_type", "Full-Time"),
        "reporting_to":     data.get("reporting_to", ""),
        "openings":         data.get("openings", ""),
        "job_code":         data.get("job_code", ""),
        "generated_by":     session.get("user", {}).get("email", ""),
        "generated_date":   datetime.now().strftime("%B %d, %Y"),
        "logo_main":        LOGO_MAIN,
        "logo_gcc":         LOGO_GCC,
        "jd_footer_text":        JD_FOOTER_TEXT,
        "jd_footer_heading":     JD_FOOTER_HEADING,
        "jd_footer_brand_label": JD_FOOTER_BRAND_LABEL,
        "jd_footer_brand_name":  JD_FOOTER_BRAND_NAME,
    }


@app.route("/api/jd/download-html", methods=["POST"])
@login_required
def api_jd_download_html():
    """
    Render the Arvind-branded HTML template and return as a downloadable .html file.
    Opening it in Chrome/Safari and using File → Print → Save as PDF gives a
    pixel-perfect branded PDF with zero server-side dependencies.
    """
    data = request.get_json(silent=True) or {}
    if not data.get("jd", "").strip():
        return jsonify({"error": "No JD content provided."}), 400

    ctx = _build_export_context(data)
    rendered = render_template("jd_export.html", **ctx)

    role_title = ctx["role_title"]
    filename = f"{role_title.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.html"
    response = make_response(rendered)
    response.headers["Content-Type"] = "text/html; charset=utf-8"
    response.headers["Content-Disposition"] = f'attachment; filename="{filename}"'
    return response


@app.route("/api/jd/preview", methods=["POST"])
@login_required
def api_jd_preview():
    """
    Render the Arvind-branded HTML in-browser (no download).
    Frontend opens this in a new tab — user can print/save-as-PDF from there.
    """
    data = request.get_json(silent=True) or {}
    if not data.get("jd", "").strip():
        return jsonify({"error": "No JD content provided."}), 400

    ctx = _build_export_context(data)
    return render_template("jd_export.html", **ctx)


@app.route("/api/jd/download-pdf", methods=["POST"])
@login_required
def api_jd_download_pdf():
    """
    Download JD as a styled PDF using reportlab.
    Matches Arvind brand: #B01030 red headers, Calibri-equiv font, section structure.
    """
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.lib.units import cm, mm
        from reportlab.lib import colors
        from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                        HRFlowable, Table, TableStyle)
        from reportlab.lib.enums import TA_LEFT, TA_RIGHT
    except ImportError:
        return jsonify({"error": "reportlab not installed. Run: pip install reportlab"}), 500

    data = request.get_json(silent=True) or {}
    jd_text    = data.get("jd", "").strip()
    role_title = data.get("role_title", "Job Description").strip()

    if not jd_text:
        return jsonify({"error": "No JD content provided."}), 400

    sections = parse_jd_sections(jd_text)

    ARVIND_RED  = colors.HexColor("#B01030")
    DARK_TEXT   = colors.HexColor("#3d3d3d")
    MID_TEXT    = colors.HexColor("#777777")
    LIGHT_BG    = colors.HexColor("#f7f7f7")
    WHITE       = colors.white

    def _style(name, **kw):
        base = dict(fontName="Helvetica", fontSize=11, textColor=DARK_TEXT,
                    leading=16, spaceAfter=4)
        base.update(kw)
        return ParagraphStyle(name, **base)

    S_TITLE   = _style("Title",   fontName="Helvetica-Bold", fontSize=20,
                        textColor=WHITE, spaceAfter=2, leading=24)
    S_META    = _style("Meta",    fontSize=9, textColor=MID_TEXT, spaceAfter=10)
    S_HEAD    = _style("Head",    fontName="Helvetica-Bold", fontSize=9,
                        textColor=ARVIND_RED, spaceBefore=16, spaceAfter=4,
                        leading=12, letterSpacing=0.8)
    S_BODY    = _style("Body",    fontSize=11, textColor=DARK_TEXT,
                        leading=17, spaceAfter=4)
    S_BULLET  = _style("Bullet",  fontSize=11, textColor=DARK_TEXT,
                        leftIndent=14, firstLineIndent=-10,
                        leading=17, spaceAfter=3)
    S_TABLE_L = _style("TblL",    fontName="Helvetica-Bold", fontSize=11,
                        textColor=DARK_TEXT, leading=15)
    S_TABLE_R = _style("TblR",    fontSize=11, textColor=DARK_TEXT, leading=15)

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            leftMargin=2*cm, rightMargin=2*cm,
                            topMargin=2.2*cm, bottomMargin=2.2*cm)

    story = []

    # ── Header banner (red) ───────────────────────────────────────────────
    banner_data = [[Paragraph(role_title, S_TITLE)]]
    banner_tbl = Table(banner_data, colWidths=[doc.width])
    banner_tbl.setStyle(TableStyle([
        ("BACKGROUND", (0,0), (-1,-1), ARVIND_RED),
        ("TOPPADDING",    (0,0), (-1,-1), 12),
        ("BOTTOMPADDING", (0,0), (-1,-1), 12),
        ("LEFTPADDING",   (0,0), (-1,-1), 14),
        ("RIGHTPADDING",  (0,0), (-1,-1), 14),
    ]))
    story.append(banner_tbl)
    story.append(Spacer(1, 4))

    # Meta line
    meta_str = (f"Generated {datetime.now().strftime('%B %d, %Y')}  ·  "
                f"{session['user']['email']}  ·  Arvind Limited")
    story.append(Paragraph(meta_str, S_META))
    story.append(HRFlowable(width="100%", thickness=1, color=ARVIND_RED,
                             spaceAfter=10, spaceBefore=0))

    # ── Position overview table ───────────────────────────────────────────
    yoe_band   = data.get("yoe_band", "")
    department = data.get("department", "").replace("_"," ").title()
    division   = data.get("division", "")
    rows = [
        ["Job Title",        role_title],
        ["Department",       department or "—"],
        ["Division / BU",    division or "—"],
        ["Location",         data.get("location", "Ahmedabad, Gujarat")],
        ["Employment Type",  data.get("employment_type", "Full-Time")],
        ["Experience Level", (yoe_band + " years") if yoe_band else "—"],
        ["Reporting To",     data.get("reporting_to", "—")],
        ["No. of Openings",  data.get("openings", "—")],
        ["Job Code",         data.get("job_code", "—")],
    ]
    col_w = [doc.width * 0.38, doc.width * 0.62]
    tbl_data = [[Paragraph(r[0], S_TABLE_L), Paragraph(r[1], S_TABLE_R)] for r in rows]
    pos_tbl = Table(tbl_data, colWidths=col_w, repeatRows=0)
    tbl_style = TableStyle([
        ("TOPPADDING",    (0,0), (-1,-1), 6),
        ("BOTTOMPADDING", (0,0), (-1,-1), 6),
        ("LEFTPADDING",   (0,0), (-1,-1), 8),
        ("RIGHTPADDING",  (0,0), (-1,-1), 8),
        ("VALIGN",        (0,0), (-1,-1), "TOP"),
    ])
    # Alternating row backgrounds
    for i in range(len(tbl_data)):
        bg = LIGHT_BG if i % 2 == 0 else WHITE
        tbl_style.add("BACKGROUND", (0,i), (0,i), bg)
    pos_tbl.setStyle(tbl_style)

    story.append(Paragraph("POSITION OVERVIEW", S_HEAD))
    story.append(pos_tbl)

    # ── Helper to add a section ───────────────────────────────────────────
    def add_section(heading: str, items):
        story.append(Paragraph(heading.upper(), S_HEAD))
        story.append(HRFlowable(width="100%", thickness=1.5, color=ARVIND_RED,
                                 spaceAfter=6, spaceBefore=0))
        if isinstance(items, str):
            for para in items.split("\n"):
                if para.strip():
                    story.append(Paragraph(para.strip(), S_BODY))
        elif isinstance(items, list):
            for item in items:
                story.append(Paragraph(f"• {item}", S_BULLET))

    # ── Sections ─────────────────────────────────────────────────────────
    if sections.get("role_summary"):
        add_section("Role Summary", sections["role_summary"])
    if sections.get("key_responsibilities"):
        add_section("Key Responsibilities", sections["key_responsibilities"])
    if sections.get("must_have"):
        add_section("Must-Have Skills & Qualifications", sections["must_have"])
    if sections.get("nice_to_have"):
        add_section("Nice-to-Have Skills", sections["nice_to_have"])

    offer = sections.get("what_we_offer") or [
        "Competitive compensation aligned to market benchmarks",
        "Opportunity to build from the ground up within a 95-year-old enterprise",
        "Exposure to global operations across Textile, Retail, Advanced Materials, and more",
        "Learning & development via Arvind's Samarth early-career platform",
        "Inclusive, diverse, and purpose-driven work culture",
    ]
    add_section("What We Offer", offer)

    # How to apply
    story.append(Paragraph("HOW TO APPLY", S_HEAD))
    story.append(HRFlowable(width="100%", thickness=1.5, color=ARVIND_RED,
                             spaceAfter=6, spaceBefore=0))
    job_code = data.get("job_code", "")
    apply_text = ("Apply through <b>careers.arvind.com</b> or email your profile to "
                  "<b>hr.ta@arvind.in</b>")
    if job_code:
        apply_text += f" with subject: <b>{job_code} — {role_title}</b>"
    story.append(Paragraph(apply_text, S_BODY))
    story.append(Spacer(1, 4))
    story.append(Paragraph(
        "<i>Arvind Limited is an equal opportunity employer committed to diversity and inclusion.</i>",
        _style("Disc", fontSize=10, textColor=MID_TEXT, leading=14)
    ))

    story.append(Spacer(1, 16))
    foot_left = (
        f"<b>{JD_FOOTER_HEADING}</b><br/>"
        f"<font size='9'>{JD_FOOTER_TEXT}</font>"
    )
    foot_right = (
        f"<font size='7' color='#ffffffaa'>{JD_FOOTER_BRAND_LABEL}</font><br/>"
        f"<b>{JD_FOOTER_BRAND_NAME}</b>"
    )
    foot_tbl = Table(
        [[Paragraph(foot_left, _style("FootL", textColor=WHITE, fontSize=11, leading=14)),
          Paragraph(foot_right, _style("FootR", textColor=WHITE, fontSize=11, leading=14, alignment=2))]],
        colWidths=[doc.width * 0.72, doc.width * 0.28],
    )
    foot_tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), ARVIND_RED),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 12),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 12),
        ("LEFTPADDING", (0, 0), (0, 0), 14),
        ("RIGHTPADDING", (-1, 0), (-1, -1), 14),
        ("LINEAFTER", (0, 0), (0, -1), 0.5, colors.HexColor("#ffffff55")),
    ]))
    story.append(foot_tbl)

    doc.build(story)
    buf.seek(0)

    filename = f"{role_title.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.pdf"
    response = make_response(buf.read())
    response.headers["Content-Type"] = "application/pdf"
    response.headers["Content-Disposition"] = f'attachment; filename="{filename}"'
    return response


@app.route("/api/download", methods=["POST"])
@login_required
def api_download():
    """Download JD as .docx (retained from Phase 1)."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    data = request.get_json(silent=True) or {}
    jd_text = strip_jd_footer(data.get("jd", "").strip())
    role_title = data.get("role_title", "Job Description").strip()

    doc = Document()
    for para in doc.paragraphs:
        para._element.getparent().remove(para._element)

    section = doc.sections[0]
    section.top_margin = Pt(72)
    section.bottom_margin = Pt(72)
    section.left_margin = Pt(80)
    section.right_margin = Pt(80)

    title_para = doc.add_heading(role_title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if title_para.runs:
        title_para.runs[0].font.size = Pt(22)
        title_para.runs[0].font.color.rgb = RGBColor(0x00, 0x37, 0x5E)

    meta = doc.add_paragraph()
    mr = meta.add_run(
        f"Generated on {datetime.now().strftime('%B %d, %Y')}  ·  {session['user']['email']}"
    )
    mr.font.size = Pt(9)
    mr.font.color.rgb = RGBColor(0x72, 0x77, 0x7F)

    doc.add_paragraph()

    SECTION_HEADINGS = {
        "role summary", "key responsibilities", "responsibilities",
        "must-have skills", "must have skills", "nice-to-have skills",
        "nice to have skills", "what we offer", "about the role",
        "the role", "requirements",
    }

    def _is_heading(line: str, prev_line: str) -> bool:
        if line.startswith(("# ", "## ", "### ")):
            return True
        if line.startswith("**") and line.endswith("**"):
            return True
        clean = line.lstrip("#").strip().rstrip(":").lower()
        if clean in SECTION_HEADINGS:
            return True
        return (
            len(line) <= 60
            and not line.startswith(("-", "•", "*", " ", "\t"))
            and line[0].isupper()
            and prev_line.strip() == ""
        )

    lines = jd_text.split("\n")
    for idx, line in enumerate(lines):
        line = line.rstrip()
        prev = lines[idx - 1] if idx > 0 else ""

        if not line:
            doc.add_paragraph()
            continue

        stripped = line.lstrip("#* ").rstrip("*").strip()

        if _is_heading(line, prev):
            h = doc.add_heading(stripped, level=2)
            if h.runs:
                h.runs[0].font.color.rgb = RGBColor(0x00, 0x37, 0x5E)
                h.runs[0].font.size = Pt(13)
        elif line.startswith(("- ", "• ", "* ")):
            p = doc.add_paragraph(line[2:], style="List Bullet")
            if p.runs:
                p.runs[0].font.size = Pt(11)
        elif len(line) > 2 and line[0].isdigit() and line[1] in ".):":
            p = doc.add_paragraph(line[2:].strip(), style="List Number")
            if p.runs:
                p.runs[0].font.size = Pt(11)
        else:
            p = doc.add_paragraph(stripped)
            if p.runs:
                p.runs[0].font.size = Pt(11)

    doc.add_paragraph()
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    def _shade(cell, fill):
        sh = OxmlElement("w:shd")
        sh.set(qn("w:fill"), fill)
        sh.set(qn("w:val"), "clear")
        cell._tc.get_or_add_tcPr().append(sh)

    ft = doc.add_table(rows=1, cols=2)
    ft.autofit = False
    ft.columns[0].width = Pt(380)
    ft.columns[1].width = Pt(120)
    c0, c1 = ft.rows[0].cells
    _shade(c0, "B01030")
    _shade(c1, "B01030")
    p0 = c0.paragraphs[0]
    r0 = p0.add_run(JD_FOOTER_HEADING + "\n")
    r0.bold = True
    r0.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    r0.font.size = Pt(12)
    r1 = p0.add_run(JD_FOOTER_TEXT)
    r1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    r1.font.size = Pt(10)
    p1 = c1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    b0 = p1.add_run(JD_FOOTER_BRAND_LABEL + "\n")
    b0.font.size = Pt(7)
    b0.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    b1 = p1.add_run(JD_FOOTER_BRAND_NAME)
    b1.bold = True
    b1.font.size = Pt(11)
    b1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = f"{role_title.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.docx"
    response = make_response(buf.read())
    response.headers["Content-Type"] = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    response.headers["Content-Disposition"] = f'attachment; filename="{filename}"'
    return response


@app.route("/api/jd/suggest-skills", methods=["GET"])
@login_required
def api_suggest_skills():
    """Return auto-suggested skills for a role title + department + family.
    First tries the hardcoded taxonomy; falls back to LLM when no results."""
    from role_skills_map import suggest_skills
    role_title = request.args.get("role_title", "").strip()
    department = request.args.get("department", "").strip()
    family     = request.args.get("family", "").strip()

    result = suggest_skills(role_title, department, family)

    # Merge in community-promoted skills for this role (dept/family), regardless
    # of whether a canonical taxonomy match was found above.
    if department and family:
        promoted = get_promoted_skills(department, family)
        existing_lower = {s.lower() for s in result.get("must_have", [])} \
            | {s.lower() for s in result.get("nice_to_have", [])}
        new_promoted = [s for s in promoted if s.lower() not in existing_lower]
        if new_promoted:
            result["must_have"] = [*result.get("must_have", []), *new_promoted]
            result["community_skills"] = new_promoted
            print(f"[COMMUNITY SKILLS SURFACED] {department}/{family}: {new_promoted}", flush=True)

    if result.get("must_have"):
        return jsonify(result)

    # Fallback: LLM-powered skill suggestion
    if not role_title and not family and not department:
        return jsonify({"must_have": [], "nice_to_have": [], "exclude": []})
    try:
        from groq import Groq
        import json as _json
        client = Groq()
        family_label = display_name(family) if family else ""
        dept_label   = DEPARTMENT_DISPLAY_NAMES.get(department, display_name(department)) if department else ""
        role_hint    = role_title or family_label or dept_label
        prompt = (
            f"List must-have technical skills for a {role_hint} role"
            + (f" in {dept_label}" if dept_label and dept_label != role_hint else "")
            + ".\nReturn JSON only — no markdown, no preamble:\n"
            '{"must_have": ["skill1", "skill2", ...], "nice_to_have": ["skill1", ...]}\n'
            "6-8 must_have (specific tools/frameworks/certs), 3-4 nice_to_have."
        )
        resp = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=300,
            temperature=0.3,
        )
        text = resp.choices[0].message.content.strip()
        start, end = text.find("{"), text.rfind("}") + 1
        if start >= 0 and end > start:
            data = _json.loads(text[start:end])
            return jsonify({
                "must_have":    data.get("must_have", []),
                "nice_to_have": data.get("nice_to_have", []),
                "exclude":      [],
                "source":       "llm",
            })
    except Exception:
        pass
    return jsonify({"must_have": [], "nice_to_have": [], "exclude": []})


@app.route("/api/jd/improve-section", methods=["POST"])
@login_required
def api_improve_section():
    """Generate AI suggestions for a specific section of a JD."""
    from groq import Groq
    data       = request.get_json(silent=True) or {}
    section    = data.get("section", "").strip()
    current    = data.get("current_text", "").strip()
    role_title = data.get("role_title", "").strip()
    department = data.get("department", "").strip()
    family     = data.get("family", "").strip()
    yoe_band   = data.get("yoe_band", "2-5").strip()
    role_type  = data.get("role_type", "individual_contributor")
    action     = data.get("action", "suggest")   # suggest | rewrite | more

    if not section:
        return jsonify({"error": "section is required"}), 400

    yoe_label    = YOE_LABELS.get(yoe_band, yoe_band)
    dept_label   = DEPARTMENT_DISPLAY_NAMES.get(department, display_name(department))
    family_label = display_name(family)
    role_label   = "Team Manager" if role_type == "team_manager" else "Individual Contributor"

    SECTION_RULES = {
        "ROLE SUMMARY": (
            "Write 3 alternative role summary paragraphs (3–4 sentences each). "
            "Each takes a different emphasis: (a) technical depth, (b) business impact, (c) team/career angle. "
            "No numbered headings — each paragraph is standalone."
        ),
        "KEY RESPONSIBILITIES": (
            "Suggest 5 additional strong, specific bullet points for key responsibilities. "
            "Each starts with a distinct strong action verb. No generic bullets. "
            "Each bullet = one real observable task."
        ),
        "MUST-HAVE SKILLS": (
            "Suggest 5 additional must-have skills/tools. "
            "Name exact tools, frameworks, languages, or certifications — not categories. "
            "Format: 'Tool/skill — what they do with it'. "
            "These must be genuinely non-negotiable for this exact role."
        ),
        "GOOD TO HAVE": (
            "Suggest 4 genuine differentiator skills — not basic requirements in disguise. "
            "Think: niche tools, adjacent domain knowledge, valuable certifications, or rare skills "
            "that make a candidate stand out for this specific role."
        ),
        "WHAT WE OFFER": (
            "Suggest 3 compelling 'What We Offer' bullet points. "
            "Focus on: real technical challenges/scale, career growth specifics, cross-functional exposure, "
            "meaningful perks relevant to this role. No clichés ('great team', 'competitive salary')."
        ),
    }

    section_key = section.upper().strip()
    instruction = SECTION_RULES.get(
        section_key,
        f"Suggest 4–5 improved alternatives for the {section} section. Be specific and concrete."
    )

    action_prefix = {
        "suggest": "Generate additional ideas to supplement the current content:",
        "rewrite": "Rewrite this section completely with a fresh perspective:",
        "more":    "Generate more alternatives — avoid repeating ideas already present:",
    }.get(action, "Improve this section:")

    prompt = f"""Context:
Role: {role_title} | Dept: {dept_label} | Family: {family_label}
Experience: {yoe_band} yrs ({yoe_label}) | Type: {role_label}

Section: {section}
{action_prefix}
---
{current}
---

{instruction}

Return ONLY the content — no section headings, no "Here are..." preamble. Each bullet starts with •"""

    try:
        client = Groq()
        resp = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[
                {"role": "system", "content": "You are a senior JD writer. Return only the requested content with no preamble or explanations."},
                {"role": "user",   "content": prompt},
            ],
            max_tokens=600,
            temperature=0.85,
        )
        raw  = resp.choices[0].message.content.strip()
        lines = [ln.lstrip("•-* \t").strip() for ln in raw.split("\n") if ln.strip()]
        return jsonify({"suggestions": lines, "raw": raw})
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/api/jd/archive", methods=["POST"])
@login_required
def api_jd_archive():
    """Archive a draft JD so it no longer appears in active history."""
    data = request.get_json(silent=True) or {}
    generation_id = data.get("generation_id", "").strip()
    if not generation_id:
        return jsonify({"error": "generation_id is required."}), 400
    sdb.archive_jd(generation_id, session["user"]["email"])
    _pending_generations.pop(generation_id, None)
    _save_pending(_pending_generations)
    return jsonify({"success": True})


@app.route("/api/jd/history", methods=["GET"])
@login_required
def api_jd_history():
    """Return approved JDs metadata for the current user, most recent first.
    Prefers Supabase; falls back to filesystem KB scan."""
    user_email = session["user"]["email"]

    # Supabase path (richer — includes drafts and all statuses)
    sb_rows = sdb.get_user_jds(user_email, status="approved")
    if sb_rows:
        results = []
        for j in sb_rows:
            results.append({
                "jd_id":       j.get("jd_ref", ""),
                "generation_id": j.get("generation_id", ""),
                "dept":        j.get("department", ""),
                "dept_label":  DEPARTMENT_DISPLAY_NAMES.get(j.get("department", ""), display_name(j.get("department", ""))),
                "family":      j.get("family", ""),
                "family_label": display_name(j.get("family", "")),
                "yoe_band":    j.get("yoe_band", ""),
                "role_title":  j.get("role_title", ""),
                "date_added":  (j.get("approved_at") or j.get("created_at", ""))[:10],
                "approved_by": j.get("user_email", ""),
                "edit_ratio":  j.get("edit_ratio"),
            })
        return jsonify({"history": results[:50], "source": "supabase"})

    # Filesystem fallback
    results = []
    for dept_dir in sorted(KB_ROOT.iterdir()):
        if not dept_dir.is_dir():
            continue
        for family_dir in sorted(dept_dir.iterdir()):
            if not family_dir.is_dir():
                continue
            sample_dir = family_dir / "sample_jds"
            if not sample_dir.exists():
                continue
            for jd_file in sample_dir.iterdir():
                if not jd_file.name.endswith(".txt"):
                    continue
                raw = jd_file.read_text(encoding="utf-8").strip()
                first_line = raw.split("\n")[0]
                try:
                    meta = json.loads(first_line)
                except Exception:
                    continue
                if meta.get("approved_by") != user_email and not DEV_MODE:
                    continue
                # approved_at (full ISO timestamp) sorts true chronological order;
                # older JDs only have date_added (day-granularity) — still works for ordering.
                sort_key = meta.get("approved_at") or meta.get("date_added", "")
                results.append({
                    "file": jd_file.name,
                    "dept": dept_dir.name,
                    "dept_label": DEPARTMENT_DISPLAY_NAMES.get(dept_dir.name, display_name(dept_dir.name)),
                    "family": family_dir.name,
                    "family_label": meta.get("role_family") or display_name(family_dir.name),
                    "yoe_band": meta.get("yoe_band", ""),
                    "date_added": meta.get("date_added", ""),
                    "approved_by": meta.get("approved_by", ""),
                    "jd_id": meta.get("jd_id", ""),
                    "_sort_key": sort_key,
                })

    # Most recently approved first
    results.sort(key=lambda r: r["_sort_key"], reverse=True)
    for r in results:
        del r["_sort_key"]

    return jsonify({"history": results[:50]})  # cap at 50


@app.route("/api/jd/draft/list", methods=["GET"])
@login_required
def api_jd_draft_list():
    """Return pending/draft JDs for the current user from Supabase."""
    user_email = session["user"]["email"]
    rows = sdb.get_user_jds(user_email, status="draft")
    drafts = [{
        "generation_id": j.get("generation_id", ""),
        "role_title":    j.get("role_title", ""),
        "department":    j.get("department", ""),
        "family":        j.get("family", ""),
        "yoe_band":      j.get("yoe_band", ""),
        "created_at":    j.get("created_at", ""),
        "jd_text":       j.get("original_text", ""),
    } for j in rows]
    return jsonify({"drafts": drafts})


@app.route("/api/admin/users", methods=["GET"])
@admin_required
def api_admin_users():
    """Return all user profiles for the admin panel."""
    profiles = sdb.get_all_user_profiles()
    return jsonify({"users": profiles})


@app.route("/api/admin/all-jds", methods=["GET"])
@admin_required
def api_admin_all_jds():
    """Return all JDs across all users (admin only)."""
    jds = sdb.get_all_jds(limit=200)
    for j in jds:
        j["dept_label"]   = DEPARTMENT_DISPLAY_NAMES.get(j.get("department", ""), display_name(j.get("department", "")))
        j["family_label"] = display_name(j.get("family", ""))
    return jsonify({"jds": jds})


@app.route("/api/jd/department-samples", methods=["GET"])
@login_required
def api_jd_department_samples():
    """
    Return one representative sample JD per department, for the History page's
    department segments. Picks the most recently added sample JD across all
    role families within each department, preferring real scraped JDs.
    """
    samples = []
    for dept, families in DEPARTMENTS.items():
        best = None  # (priority, date_added, payload)
        for family in families:
            sample_dir = KB_ROOT / dept / family / "sample_jds"
            if not sample_dir.exists():
                continue
            for jd_file in sorted(sample_dir.iterdir(), reverse=True):
                if not jd_file.name.endswith(".txt"):
                    continue
                raw = jd_file.read_text(encoding="utf-8").strip()
                if "\n\n---\n\n" not in raw:
                    continue
                meta_str, body = raw.split("\n\n---\n\n", 1)
                try:
                    meta = json.loads(meta_str)
                except Exception:
                    continue

                # Prefer real scraped JDs > user-approved > KB seed
                source = meta.get("source", "")
                priority = {"real_scraped_jd": 2, "kb_seed": 0}.get(source, 1)
                date_added = meta.get("date_added", "")

                candidate = (priority, date_added, {
                    "dept": dept,
                    "dept_label": DEPARTMENT_DISPLAY_NAMES.get(dept, display_name(dept)),
                    "family": family,
                    "family_label": display_name(family),
                    "role_title": meta.get("role_family", display_name(family)),
                    "yoe_band": meta.get("yoe_band", ""),
                    "preview": body[:280].strip() + ("…" if len(body) > 280 else ""),
                    "full_text": body.strip(),
                    "jd_id": meta.get("jd_id", ""),
                })

                if best is None or candidate[:2] > best[:2]:
                    best = candidate

        if best is not None:
            samples.append(best[2])
        else:
            samples.append({
                "dept": dept,
                "dept_label": DEPARTMENT_DISPLAY_NAMES.get(dept, display_name(dept)),
                "family": families[0] if families else "",
                "family_label": display_name(families[0]) if families else "",
                "role_title": "No sample available yet",
                "yoe_band": "",
                "preview": "Generate and approve a JD for this department to populate this card.",
                "full_text": "",
                "jd_id": "",
            })

    return jsonify({"samples": samples})


# ── Dev helper ────────────────────────────────────────────────────────────────

@app.route("/api/departments", methods=["GET"])
def api_departments():
    """Return departments + families for frontend dropdowns, with canonical display labels."""
    return jsonify({
        "departments": DEPARTMENTS,
        "department_labels": DEPARTMENT_DISPLAY_NAMES,
        "divisions": DIVISIONS,
        "yoe_bands": YOE_BANDS,
        "yoe_labels": YOE_LABELS,
    })


# ── Sample JD Library (new — zero-LLM-cost browsing mode) ────────────────────
# Reads directly off the kb/ filesystem so it never depends on the casing or
# contents of the DEPARTMENTS/DIVISIONS dicts used elsewhere in the app.

import base64
import difflib


def _division_for_dept(dept_folder: str) -> str:
    """Best-effort lookup of which Division a department folder belongs to."""
    dept_lower = dept_folder.lower()
    for division_name, dept_list in DIVISIONS.items():
        if any(d.lower() == dept_lower for d in dept_list):
            return division_name
    return "Other"


def _make_sample_ref(dept: str, family: str, filename: str) -> str:
    """Build a URL-safe opaque reference for a sample JD file."""
    raw = f"{dept}/{family}/{filename}"
    return base64.urlsafe_b64encode(raw.encode("utf-8")).decode("ascii").rstrip("=")


def _resolve_sample_ref(ref: str) -> Path | None:
    """Reverse _make_sample_ref back into a real file path, with a path-traversal guard."""
    padded = ref + "=" * (-len(ref) % 4)
    try:
        raw = base64.urlsafe_b64decode(padded.encode("ascii")).decode("utf-8")
    except Exception:
        return None
    parts = raw.split("/")
    if len(parts) != 3 or any(p in ("", ".", "..") for p in parts):
        return None
    candidate = (KB_ROOT / parts[0] / parts[1] / "sample_jds" / parts[2]).resolve()
    if KB_ROOT.resolve() not in candidate.parents:
        return None
    return candidate if candidate.exists() else None


def _iter_sample_files():
    """Yield (dept_folder, family_folder, file_path) for every sample JD on disk."""
    for dept_dir in sorted(KB_ROOT.iterdir()):
        if not dept_dir.is_dir():
            continue
        for family_dir in sorted(dept_dir.iterdir()):
            if not family_dir.is_dir():
                continue
            sample_dir = family_dir / "sample_jds"
            if not sample_dir.exists():
                continue
            for jd_file in sorted(sample_dir.iterdir()):
                if jd_file.name.endswith(".txt"):
                    yield dept_dir.name, family_dir.name, jd_file


def _load_sample_file(jd_file: Path) -> tuple[dict, str] | None:
    raw = jd_file.read_text(encoding="utf-8").strip()
    if "\n\n---\n\n" not in raw:
        return None
    meta_str, body = raw.split("\n\n---\n\n", 1)
    try:
        meta = json.loads(meta_str)
    except Exception:
        meta = {}
    return meta, body.strip()


@app.route("/library")
@login_required
def library():
    return render_template("library.html", user=_enrich_user(session["user"]))


@app.route("/api/sample/filters", methods=["GET"])
@login_required
def api_sample_filters():
    """Division → Department → Role Family options, derived live from kb/ on disk."""
    dept_families: dict[str, list[str]] = {}
    for dept_dir in sorted(KB_ROOT.iterdir()):
        if not dept_dir.is_dir():
            continue
        families = [f.name for f in sorted(dept_dir.iterdir()) if f.is_dir()]
        if families:
            dept_families[dept_dir.name] = families

    divisions: dict[str, list[str]] = {}
    for dept in dept_families:
        division = _division_for_dept(dept)
        divisions.setdefault(division, []).append(dept)

    return jsonify({
        "divisions": divisions,
        "dept_families": dept_families,
        "dept_labels": {d: display_name(d) for d in dept_families},
        "family_labels": {f: display_name(f) for fams in dept_families.values() for f in fams},
        "yoe_bands": YOE_BANDS,
        "yoe_labels": YOE_LABELS,
    })


@app.route("/api/sample/list", methods=["GET"])
@login_required
def api_sample_list():
    """List approved JDs for the Sample JD Library.
    Supabase-approved JDs come first (newest → oldest), then KB seed files as fallback."""
    f_dept     = request.args.get("dept",     "").strip().lower()
    f_family   = request.args.get("family",   "").strip().lower()
    f_yoe      = request.args.get("yoe_band", "").strip()
    f_division = request.args.get("division", "").strip()
    f_search   = request.args.get("search",   "").strip().lower()

    results = []

    # ── 1. Supabase: all approved JDs (newest first) ───────────────────────
    sb_rows = sdb.get_all_approved_jds(limit=300)
    for row in sb_rows:
        dept   = (row.get("department") or "").lower()
        family = (row.get("family")     or "").lower()
        yoe    = row.get("yoe_band", "")
        division = _division_for_dept(dept)
        role_title = row.get("role_title") or display_name(family)

        if f_dept     and dept     != f_dept:     continue
        if f_family   and family   != f_family:   continue
        if f_yoe      and yoe      != f_yoe:      continue
        if f_division and division != f_division: continue
        if f_search   and f_search not in role_title.lower(): continue

        # Use generation_id as the ref (prefixed so api_sample_get knows it's Supabase)
        gen_id  = row.get("generation_id", "")
        approved_at = (row.get("approved_at") or row.get("created_at") or "")[:10]

        results.append({
            "ref":          f"sb:{gen_id}",
            "role_title":   role_title,
            "dept":         dept,
            "dept_label":   DEPARTMENT_DISPLAY_NAMES.get(dept, display_name(dept)),
            "division":     division,
            "family":       family,
            "family_label": display_name(family),
            "yoe_band":     yoe,
            "date_added":   approved_at,
            "approved_by":  row.get("user_email", ""),
            "source":       "approved",
        })

    # ── 2. KB seed files as fallback (only if no Supabase results at all) ──
    if not results:
        for dept, family, jd_file in _iter_sample_files():
            if f_dept   and dept.lower()   != f_dept:   continue
            if f_family and family.lower() != f_family: continue

            division = _division_for_dept(dept)
            if f_division and division != f_division:
                continue

            loaded = _load_sample_file(jd_file)
            if not loaded:
                continue
            meta, body = loaded
            if not body.strip():
                continue

            yoe_band = meta.get("yoe_band", "")
            if f_yoe and yoe_band != f_yoe:
                continue

            role_title = meta.get("role_family", display_name(family))
            if f_search and f_search not in role_title.lower():
                continue

            results.append({
                "ref":          _make_sample_ref(dept, family, jd_file.name),
                "role_title":   role_title,
                "dept":         dept,
                "dept_label":   display_name(dept),
                "division":     division,
                "family":       family,
                "family_label": display_name(family),
                "yoe_band":     yoe_band,
                "date_added":   meta.get("date_added", ""),
                "source":       "kb",
            })
        results.sort(key=lambda r: r.get("date_added", "") or "", reverse=True)

    return jsonify({"jds": results})


@app.route("/api/sample/get", methods=["GET"])
@login_required
def api_sample_get():
    """Full JD content + parsed sections for the preview modal.
    Handles both Supabase-approved refs (sb:<generation_id>) and KB file refs."""
    ref = (request.args.get("jd_id") or request.args.get("ref") or "").strip()

    # ── Supabase-approved JD ───────────────────────────────────────────────
    if ref.startswith("sb:"):
        gen_id = ref[3:]
        row = sdb.get_jd_by_generation_id(gen_id)
        if not row:
            return jsonify({"error": "Approved JD not found."}), 404

        body   = (row.get("final_text") or row.get("original_text") or "").strip()
        dept   = (row.get("department") or "").lower()
        family = (row.get("family")     or "").lower()
        sections = parse_jd_sections(body)
        approved_at = (row.get("approved_at") or row.get("created_at") or "")[:10]

        return jsonify({
            "jd": {
                "ref":          ref,
                "role_title":   row.get("role_title") or display_name(family),
                "dept":         dept,
                "dept_label":   DEPARTMENT_DISPLAY_NAMES.get(dept, display_name(dept)),
                "division":     _division_for_dept(dept),
                "family":       family,
                "family_label": display_name(family),
                "yoe_band":     row.get("yoe_band", ""),
                "date_added":   approved_at,
                "approved_by":  row.get("user_email", ""),
                "full_text":    body,
                **sections,
            }
        })

    # ── KB filesystem JD ───────────────────────────────────────────────────
    jd_file = _resolve_sample_ref(ref)
    if not jd_file:
        return jsonify({"error": "Sample JD not found."}), 404

    loaded = _load_sample_file(jd_file)
    if not loaded:
        return jsonify({"error": "Sample JD could not be parsed."}), 500
    meta, body = loaded

    dept   = jd_file.parent.parent.parent.name
    family = jd_file.parent.parent.name
    sections = parse_jd_sections(body)

    return jsonify({
        "jd": {
            "ref":          ref,
            "role_title":   meta.get("role_family", display_name(family)),
            "dept":         dept,
            "dept_label":   display_name(dept),
            "division":     _division_for_dept(dept),
            "family":       family,
            "family_label": display_name(family),
            "yoe_band":     meta.get("yoe_band", ""),
            "date_added":   meta.get("date_added", ""),
            "full_text":    body,
            **sections,
        }
    })


def _safe_filename_part(value: str, fallback: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9_-]+", "_", (value or fallback).replace(" ", "_"))
    return re.sub(r"_+", "_", cleaned).strip("_") or fallback


def _sample_copy_filename(role_title: str, dept: str, yoe_band: str, ext: str) -> str:
    safe_role = _safe_filename_part(role_title, "Job_Description")
    safe_dept = _safe_filename_part(dept, "Dept")
    safe_yoe = _safe_filename_part(yoe_band, "NA")
    today = datetime.now().strftime("%Y%m%d")
    return f"COPY_{safe_role}_{safe_dept}_{safe_yoe}_{today}.{ext}"


@app.route("/api/sample/export-pdf", methods=["POST"])
@login_required
def api_sample_export_pdf():
    """
    Export an in-browser-edited copy of a sample JD as PDF.
    Nothing here is written to the KB, History, or any database —
    this mirrors the 'Edit a Copy' flow which is intentionally ephemeral.
    """
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    except ImportError:
        return jsonify({"error": "reportlab not installed. Run: pip install reportlab"}), 500

    data = request.get_json(silent=True) or {}
    jd_text = data.get("jd_text", "").strip()
    role_title = data.get("role_title", "Job Description").strip()
    department = data.get("department", "")
    division = data.get("division", "")
    yoe_band = data.get("yoe_band", "")

    if not jd_text:
        return jsonify({"error": "No JD content provided."}), 400

    ARVIND_BLUE = colors.HexColor("#00375E")
    MUTED = colors.HexColor("#72777F")

    def _style(name, **kw):
        base = dict(fontName="Helvetica", fontSize=11, textColor=colors.black, leading=16, spaceAfter=4)
        base.update(kw)
        return ParagraphStyle(name, **base)

    title_style = _style("CopyTitle", fontSize=20, textColor=ARVIND_BLUE, fontName="Helvetica-Bold", spaceAfter=2)
    sub_style   = _style("CopySub", fontSize=10, textColor=MUTED, spaceAfter=10)
    head_style  = _style("CopyHead", fontSize=12, textColor=ARVIND_BLUE, fontName="Helvetica-Bold", spaceBefore=12, spaceAfter=4)
    body_style  = _style("CopyBody", fontSize=11)
    foot_style  = _style("CopyFoot", fontSize=9, textColor=MUTED)

    sections = parse_jd_sections(jd_text)

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=2.2*cm, rightMargin=2.2*cm,
                            topMargin=2.2*cm, bottomMargin=2.2*cm)
    story = [
        Paragraph(role_title, title_style),
        Paragraph(" · ".join(filter(None, [department, division, (yoe_band + " yrs" if yoe_band else "")])), sub_style),
        HRFlowable(width="100%", thickness=1, color=ARVIND_BLUE, spaceAfter=10),
    ]

    def add_section(heading, content):
        story.append(Paragraph(heading, head_style))
        if isinstance(content, str) and content:
            for line in content.split("\n"):
                if line.strip():
                    story.append(Paragraph(line.strip(), body_style))
        elif isinstance(content, list):
            for item in content:
                story.append(Paragraph(f"• {item}", body_style))

    if sections.get("role_summary"):
        add_section("Role Summary", sections["role_summary"])
    if sections.get("key_responsibilities"):
        add_section("Key Responsibilities", sections["key_responsibilities"])
    if sections.get("must_have"):
        add_section("Must-Have Skills", sections["must_have"])
    if sections.get("nice_to_have"):
        add_section("Nice-to-Have Skills", sections["nice_to_have"])
    if sections.get("what_we_offer"):
        add_section("What We Offer", sections["what_we_offer"])

    story.append(Spacer(1, 16))
    story.append(HRFlowable(width="100%", thickness=0.5, color=MUTED, spaceAfter=6))
    story.append(Paragraph(
        f"Arvind Limited · Copied from Sample Library · {datetime.now().strftime('%Y-%m-%d')} · For Internal Use",
        foot_style,
    ))

    doc.build(story)
    buf.seek(0)

    filename = _sample_copy_filename(role_title, department, yoe_band, "pdf")
    response = make_response(buf.read())
    response.headers["Content-Type"] = "application/pdf"
    response.headers["Content-Disposition"] = f'attachment; filename="{filename}"'
    return response


@app.route("/api/sample/export-docx", methods=["POST"])
@login_required
def api_sample_export_docx():
    """Export an in-browser-edited copy of a sample JD as DOCX. Never written to the KB."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    data = request.get_json(silent=True) or {}
    jd_text = data.get("jd_text", "").strip()
    role_title = data.get("role_title", "Job Description").strip()
    department = data.get("department", "")
    division = data.get("division", "")
    yoe_band = data.get("yoe_band", "")

    if not jd_text:
        return jsonify({"error": "No JD content provided."}), 400

    sections = parse_jd_sections(jd_text)

    doc = Document()
    for para in doc.paragraphs:
        para._element.getparent().remove(para._element)

    title = doc.add_heading(role_title, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if title.runs:
        title.runs[0].font.size = Pt(20)
        title.runs[0].font.color.rgb = RGBColor(0x00, 0x37, 0x5E)

    sub = doc.add_paragraph()
    sub_run = sub.add_run(" · ".join(filter(None, [department, division, (yoe_band + " yrs" if yoe_band else "")])))
    sub_run.font.size = Pt(9)
    sub_run.font.color.rgb = RGBColor(0x72, 0x77, 0x7F)
    sub_run.italic = True

    doc.add_paragraph()

    def add_section(heading, content):
        h = doc.add_heading(heading, level=2)
        if h.runs:
            h.runs[0].font.color.rgb = RGBColor(0x00, 0x37, 0x5E)
            h.runs[0].font.size = Pt(13)
        if isinstance(content, str) and content:
            for line in content.split("\n"):
                if line.strip():
                    p = doc.add_paragraph(line.strip())
                    if p.runs:
                        p.runs[0].font.size = Pt(11)
        elif isinstance(content, list):
            for item in content:
                p = doc.add_paragraph(item, style="List Bullet")
                if p.runs:
                    p.runs[0].font.size = Pt(11)

    if sections.get("role_summary"):
        add_section("Role Summary", sections["role_summary"])
    if sections.get("key_responsibilities"):
        add_section("Key Responsibilities", sections["key_responsibilities"])
    if sections.get("must_have"):
        add_section("Must-Have Skills", sections["must_have"])
    if sections.get("nice_to_have"):
        add_section("Nice-to-Have Skills", sections["nice_to_have"])
    if sections.get("what_we_offer"):
        add_section("What We Offer", sections["what_we_offer"])

    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer_run = footer.add_run(
        f"Arvind Limited · Copied from Sample Library · {datetime.now().strftime('%Y-%m-%d')} · Confidential"
    )
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(0x72, 0x77, 0x7F)
    footer_run.italic = True

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = _sample_copy_filename(role_title, department, yoe_band, "docx")
    response = make_response(buf.read())
    response.headers["Content-Type"] = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    response.headers["Content-Disposition"] = f'attachment; filename="{filename}"'
    return response


@app.route("/api/jd/share", methods=["POST"])
@login_required
def api_jd_share():
    """
    Placeholder share-link endpoint. Non-functional until Microsoft SSO
    credentials are configured — returns a clear message instead of a link.
    """
    return jsonify({
        "message": "Share coming soon — configure Microsoft SSO credentials to enable sharing.",
        "url": None,
    })


@app.route("/shared/<share_token>")
def view_shared_jd(share_token):
    """Placeholder shared-JD viewer. Will validate an SSO session once wired up."""
    return jsonify({
        "error": "Sharing is not yet configured for this workspace.",
        "share_token": share_token,
    }), 501


# ── Darwinbox / Job Board Distribution routes ─────────────────────────────────

@app.route("/api/darwin/jobs", methods=["GET"])
@login_required
def api_darwin_jobs():
    """Active Darwinbox JD library — filtered view data."""
    jobs = darwin.fetch_darwinbox_jobs(active_only=True)
    # Attach postings to each job
    for job in jobs:
        job["postings"] = darwin.fetch_job_postings(job["darwinbox_job_id"])
    return jsonify(jobs)


@app.route("/api/darwin/publish", methods=["POST"])
@login_required
def api_darwin_publish():
    """Publish a Darwinbox job to LinkedIn, Naukri, or both."""
    data = request.get_json(silent=True) or {}
    darwinbox_job_id = data.get("darwinbox_job_id", "").strip()
    platforms = data.get("platforms", [])  # ["linkedin"], ["naukri"], or both

    if not darwinbox_job_id or not platforms:
        return jsonify({"error": "darwinbox_job_id and platforms required."}), 400

    valid = {"linkedin", "naukri"}
    platforms = [p for p in platforms if p in valid]
    if not platforms:
        return jsonify({"error": "platforms must include linkedin, naukri, or both."}), 400

    base_url = request.host_url.rstrip("/")
    postings = []
    for platform in platforms:
        posting = darwin.publish_to_platform(darwinbox_job_id, platform, base_url)
        postings.append(posting)

    return jsonify({"success": True, "postings": postings})


@app.route("/api/darwin/candidates", methods=["GET"])
@login_required
def api_darwin_candidates():
    """Candidate library — filterable by job_id and platform_source."""
    job_id = request.args.get("job_id", "").strip() or None
    platform = request.args.get("platform", "").strip() or None
    candidates = darwin.fetch_candidates(darwinbox_job_id=job_id, platform_source=platform)
    return jsonify(candidates)


@app.route("/apply/<platform>/<darwinbox_job_id>", methods=["GET", "POST"])
def apply_page(platform, darwinbox_job_id):
    """Public-facing candidate application page — no auth required."""
    if platform not in ("linkedin", "naukri"):
        return "Invalid platform.", 404

    job = darwin.get_darwinbox_job(darwinbox_job_id)
    if not job:
        return "This job posting is no longer available.", 404

    if request.method == "GET":
        return render_template(
            "apply.html",
            job=job,
            platform=platform,
            darwinbox_job_id=darwinbox_job_id,
        )

    # POST — handle application submission
    name = request.form.get("name", "").strip()
    email = request.form.get("email", "").strip()
    phone = request.form.get("phone", "").strip()
    consent_given = request.form.get("consent") in ("1", "on", "true", "yes")

    if not all([name, email]):
        return render_template(
            "apply.html",
            job=job,
            platform=platform,
            darwinbox_job_id=darwinbox_job_id,
            error="Name and email are required.",
        )

    if not consent_given:
        return render_template(
            "apply.html",
            job=job,
            platform=platform,
            darwinbox_job_id=darwinbox_job_id,
            error="You must provide consent before submitting your application.",
        )

    resume_path = ""
    resume_file = request.files.get("resume")
    if resume_file and resume_file.filename:
        upload_dir = Path(__file__).parent / "static" / "uploads" / "resumes"
        upload_dir.mkdir(parents=True, exist_ok=True)
        safe_name = f"{uuid.uuid4()}_{resume_file.filename.replace(' ', '_')}"
        resume_file.save(upload_dir / safe_name)
        resume_path = f"uploads/resumes/{safe_name}"

    from datetime import datetime as _dt
    consent_timestamp = _dt.now().isoformat()
    candidate = darwin.submit_candidate(
        darwinbox_job_id=darwinbox_job_id,
        platform_source=platform,
        name=name,
        email=email,
        phone=phone,
        resume_file=resume_path,
        consent_given=True,
        consent_timestamp=consent_timestamp,
    )

    # Fire-and-forget: score candidate asynchronously so submission isn't blocked
    threading.Thread(
        target=scoring.scoreCandidateOnSubmit,
        args=(candidate["candidate_id"],),
        daemon=True,
    ).start()

    return render_template(
        "apply.html",
        job=job,
        platform=platform,
        darwinbox_job_id=darwinbox_job_id,
        submitted=True,
    )


@app.route("/api/candidates/score-pending", methods=["GET", "OPTIONS"])
def api_score_pending():
    """Score all candidates that don't have a score yet. Fire-and-forget per candidate."""
    origin = request.headers.get("Origin", "")
    cors = {
        "Access-Control-Allow-Origin": origin or "*",
        "Access-Control-Allow-Methods": "GET, OPTIONS",
        "Access-Control-Allow-Headers": "Content-Type",
    }
    if request.method == "OPTIONS":
        return ("", 204, cors)

    from scoring_service import _load, _CANDIDATES_FILE, get_score
    candidates = _load(_CANDIDATES_FILE)
    pending = [c for c in candidates if not get_score(c["candidate_id"], c["darwinbox_job_id"])]

    for c in pending:
        threading.Thread(
            target=scoring.scoreCandidateOnSubmit,
            args=(c["candidate_id"],),
            daemon=True,
        ).start()

    resp = jsonify({"queued": len(pending), "total": len(candidates)})
    for k, v in cors.items():
        resp.headers[k] = v
    return resp


@app.route("/api/candidates/rescore", methods=["POST", "OPTIONS"])
def api_rescore_candidate():
    # Allow cross-origin calls from DummyDarwin (localhost:5002)
    origin = request.headers.get("Origin", "")
    headers = {
        "Access-Control-Allow-Origin": origin or "*",
        "Access-Control-Allow-Methods": "POST, OPTIONS",
        "Access-Control-Allow-Headers": "Content-Type",
        "Access-Control-Allow-Credentials": "true",
    }
    if request.method == "OPTIONS":
        return ("", 204, headers)

    data = request.get_json(silent=True) or {}
    candidate_id = data.get("candidate_id")
    darwinbox_job_id = data.get("darwinbox_job_id")
    if not candidate_id or not darwinbox_job_id:
        resp = jsonify({"success": False, "error": "candidate_id and darwinbox_job_id required"})
        resp.status_code = 400
        for k, v in headers.items():
            resp.headers[k] = v
        return resp

    result = scoring.rescoreCandidate(candidate_id, darwinbox_job_id)
    resp = jsonify({"success": True, "score": result} if result else
                   {"success": False, "error": "Scoring failed — check GROQ_API_KEY and resume file"})
    if not result:
        resp.status_code = 500
    for k, v in headers.items():
        resp.headers[k] = v
    return resp


if __name__ == "__main__":
    port = int(os.getenv("PORT", 5001))
    debug = os.getenv("FLASK_ENV", "development") == "development"
    if DEV_MODE:
        print("⚠  No SUPABASE_URL found — running in dev mode (auth bypassed).")
    # use_reloader=False prevents the reloader from spawning a child process
    # that holds the port — avoids "Address already in use" on restart
    app.run(debug=debug, port=port, use_reloader=False)
